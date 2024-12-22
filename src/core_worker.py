# core_worker.py

# Import necessary modules
import os
import shutil
import logging
import zipfile
import subprocess
import hashlib  # For calculating SHA-256 hash
import sqlite3
import json
import pytz
import requests
import gettext  # For translations
from datetime import datetime
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

# Set up logging
logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# Import LangChain and OpenAI
import openai
import langchain
from langchain.cache import SQLiteCache
from langchain.chat_models import ChatOpenAI
from langchain.memory import ConversationBufferMemory
from langchain.chains import LLMChain
from langchain.prompts import (
    ChatPromptTemplate,
    SystemMessagePromptTemplate,
    HumanMessagePromptTemplate,
    MessagesPlaceholder,
)

# Set up caching
langchain.llm_cache = SQLiteCache(database_path="/app/dbs/langchain.db")

#######################################
### HELPER: DESCRIBE JWLIBRARY FILE ###
#######################################

def describe_jwlibrary(telegram_user):
    logger.info("describe_jwlibrary - Telegram User: {0}".format(telegram_user))
    jwfile = "userBackups/{0}.jwlibrary".format(telegram_user)

    with zipfile.ZipFile(jwfile, 'r') as zip_ref:
        files = zip_ref.namelist()
        zip_ref.extractall("userBackups/{0}/".format(telegram_user))

    uploadedDb = "userBackups/{0}/{1}".format(telegram_user, [zipname for zipname in files if zipname.endswith(".db")][0])

    connection = sqlite3.connect(uploadedDb)
    cursor = connection.cursor()
    cursor.execute("SELECT Count(*) FROM Note")
    notesN = cursor.fetchone()[0]
    cursor.execute("SELECT Count(*) FROM InputField")
    inputN = cursor.fetchone()[0]
    cursor.execute("SELECT Count(*) FROM TagMap")
    tagMaptN = cursor.fetchone()[0]
    cursor.execute("SELECT Count(*) FROM Tag")
    tagN = cursor.fetchone()[0]
    cursor.execute("SELECT Count(*) FROM Bookmark")
    bookmarkN = cursor.fetchone()[0]
    cursor.execute("SELECT LastModified FROM LastModified")
    lastModified = cursor.fetchone()[0]
    cursor.execute("SELECT Count(*) FROM UserMark")
    userMarkN = cursor.fetchone()[0]
    connection.close()

    shutil.rmtree("userBackups/{0}/".format(telegram_user))

    return notesN, inputN, tagMaptN, tagN, bookmarkN, lastModified, userMarkN

#######################
### EXTRACTING HTML ###
#######################

def w_extract_html(url, get_all):
    logger.info("w_extract_html - URL: {0} - Full Run: {1}".format(url, get_all))

    html = requests.get(url).text
    soup = BeautifulSoup(html, features="html5lib")
    title = soup.find("h1").text
    classArticleId = soup.find("article", {"id": "article"}).get("class")
    articleId = next(x for x in classArticleId if x.startswith("iss"))[4:] + "00"
    articleN = soup.find("p", {"id": "p1"}).text

    if get_all:
        base_text = soup.find("p", {"id": "p4"}).text
        song = soup.find("p", {"id": "p2"}).text
        summary = soup.find("p", {"id": "p6"}).text
        documentId = soup.find("input", {"name": "docid"}).get("value")
        p_elements = soup.find("div", {"class": "bodyTxt"})
        questions = p_elements.find_all("p", {"class": lambda x: x and x.startswith("qu")})
        paragraphs = p_elements.find_all("p", {"class": lambda x: x and x.startswith("p")})
        textareas = soup.find_all("textarea")

        # Example q_map = {0: [q1, [p1]], 1: [q2&3, [p2, p3]]}
        q_map = {}
        i = 0
        for q in questions:
            q_map[i] = [q]
            q_map[i].append([p for p in paragraphs if p.has_attr('data-rel-pid') and p.get('data-rel-pid').strip('[]') in q.get('data-pid')])
            i += 1

        return title, base_text, song, summary, questions, documentId, articleId, q_map, textareas
    else:
        return title, articleId, articleN

####################
### QUERY OPENAI ###
####################

def w_query_openai(title, base_text, song, summary, q_map, qs_user, trans):
    logger.info("w_query_openai - Title: {0}".format(title))

    questions = [f"{i+1}. {question}" for i, question in enumerate(qs_user) if question]
    questions_text = "\n".join(questions)
    logger.info("w_query_openai - Questions Text: {0}".format(questions_text))

    # Set up the ChatOpenAI LLM
    llm = ChatOpenAI(model_name="gpt-3.5-turbo", temperature=0)

    # Set up the prompt
    system_prompt_template = trans("""Eres un asistente que únicamente usa jw.org y las publicaciones de los Testigos de Jehová para mejorar la preparación de reuniones.
Yo estoy preparándome la Atalaya, edición de estudio, de los Testigos de Jehová.
Proveerás información extra proveniente de la literatura disponible en cada uno de los párrafos que te voy a ir mandando en los sucesivos prompts.
La Atalaya de esta semana se titula {title}, se basa en el texto de {base_text}, cantaremos la '{song}', y el resumen es el siguiente:
{summary}
Para cada pregunta y párrafo o párrafos que te vaya enviando a partir de ahora, responderás en una lista lo siguiente:
{questions_text}
No escribas estas preguntas de nuevo en la respuesta. Separa las respuestas con dos retornos de carro.""")

    system_prompt = system_prompt_template.format(
        title=title,
        base_text=base_text,
        song=song,
        summary=summary,
        questions_text=questions_text
    )

    # Log the full system prompt
    logger.info("w_query_openai - System Prompt:\n{0}".format(system_prompt))

    prompt = ChatPromptTemplate.from_messages([
        SystemMessagePromptTemplate.from_template(system_prompt),
        MessagesPlaceholder(variable_name="history"),
        HumanMessagePromptTemplate.from_template("{input}")
    ])

    notes = {}
    i = 0

    for idx, q in enumerate(q_map.values()):
        memory = ConversationBufferMemory(memory_key="history", return_messages=True)
        chain = LLMChain(llm=llm, prompt=prompt, memory=memory)

        # Flatten the paragraphs
        flattened_paragraph = "".join([p.text for p in q[1]])

        # Prepare the user input
        user_input_template = trans("Pregunta: {question} -- Párrafo(s): {paragraphs}")
        user_input = user_input_template.format(
            question=q[0].text,
            paragraphs=flattened_paragraph
        )

        # Log the user input
        logger.info("w_query_openai - User Input for question {0}:\n{1}".format(idx+1, user_input))

        # Log the full prompt (system prompt + user input)
        logger.info("w_query_openai - Full Prompt for question {0}:\n{1}\n{2}".format(
            idx+1, system_prompt, user_input))

        # Call the chain to get the response
        notes[i] = chain.predict(input=user_input)

        # Log the response
        logger.info("w_query_openai(Note) - Note for question {0}:\n{1}".format(idx+1, notes[i]))

        i += 1

    return notes

############################
### WRITE JWLIBRARY FILE ###
############################

def calculate_user_data_hash(user_data_db_path):
    sha256_hash = hashlib.sha256()
    with open(user_data_db_path, "rb") as f:
        for byte_block in iter(lambda: f.read(4096), b""):
            sha256_hash.update(byte_block)
    hash_digest = sha256_hash.hexdigest()
    return hash_digest

def get_last_modified_date(user_data_db_path):
    last_modified_timestamp = os.path.getmtime(user_data_db_path)
    last_modified_date = datetime.fromtimestamp(
        last_modified_timestamp,
        pytz.timezone('Europe/Madrid')
    ).isoformat()
    return last_modified_date

def write_jwlibrary(documentId, articleId, title, questions, notes, telegram_user, textareas):
    logger.info("write_jwlibrary - Document ID: {0} - Article ID: {1} - Title: {2}".format(documentId, articleId, title))
    uploadedJwLibrary = 'userBackups/{0}.jwlibrary'.format(telegram_user)

    os.makedirs("/app/userBackups/{0}".format(telegram_user), exist_ok=True)

    now = datetime.now(pytz.timezone('Europe/Madrid'))
    now_date = now.strftime("%Y-%m-%d")
    now_iso = now.isoformat("T", "seconds")
    now_utc = now.astimezone(pytz.UTC)
    now_utc_iso = now_utc.isoformat("T", "seconds").replace('+00:00', 'Z')
    schema_version = 14  # TODO: Upgrade when needed

    thumbnail_file = "extra/default_thumbnail.png"

    if os.path.isfile(uploadedJwLibrary):
        logger.info("Archivo .jwlibrary encontrado")
        with zipfile.ZipFile(uploadedJwLibrary, 'r') as zip_ref:
            files = zip_ref.namelist()
            zip_ref.extractall("userBackups/{0}/".format(telegram_user))

        uploadedDb = "userBackups/{0}/{1}".format(telegram_user, [zipname for zipname in files if zipname.endswith(".db")][0])
        manifestUser = "userBackups/{0}/manifest.json".format(telegram_user)

        connection = sqlite3.connect(uploadedDb)
        cursor = connection.cursor()
        cursor.execute("SELECT LocationId FROM Location WHERE DocumentId=?", (documentId,))
        locationId = cursor.fetchone()
        if locationId:
            locationId = locationId[0]
        else:
            cursor.execute("SELECT max(LocationId) FROM Location")
            max_location_id = cursor.fetchone()[0]
            locationId = max_location_id + 1 if max_location_id else 1
            cursor.execute("""INSERT INTO Location (LocationId, DocumentId, IssueTagNumber, KeySymbol, Type)
                VALUES (?, ?, ?, "w", 0);""", (locationId, documentId, articleId))

        for i in notes:
            cursor.execute("""INSERT INTO InputField ('LocationId', 'TextTag', 'Value') VALUES (?, ?, ?)""",
                           (locationId, textareas[i].get("id"), notes[i].replace("'", '"')))

        cursor.execute("UPDATE LastModified SET LastModified = ?", (now_iso,))

        connection.commit()
        connection.close()

        # Calculate hash and last modified date
        hash_digest = calculate_user_data_hash(uploadedDb)
        last_modified_date = get_last_modified_date(uploadedDb)

        # Create manifest data
        manifest_data = {
            "type": 0,
            "name": f"jwlibrary-plus-backup_{now_date}",
            "userDataBackup": {
                "deviceName": "jwlibraryPlus",
                "hash": hash_digest,
                "lastModifiedDate": last_modified_date,
                "databaseName": "userData.db",
                "schemaVersion": schema_version
            },
            "version": 1,
            "creationDate": now.isoformat()
        }

        manifest_file = 'userBackups/{0}/manifest.json'.format(telegram_user)
        with open(manifest_file, 'w', encoding='utf-8') as f:
            json.dump(manifest_data, f, ensure_ascii=False)

        fileName = "userBackups/{0}/jwlibrary-plus-{1}-{2}.jwlibrary".format(telegram_user, documentId, now_date)
        with zipfile.ZipFile(fileName, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.write(uploadedDb, arcname="userData.db")
            zf.write(manifest_file, arcname="manifest.json")
            zf.write(thumbnail_file, arcname="default_thumbnail.png")

        os.remove(uploadedDb)
        os.remove(manifest_file)
        os.remove(uploadedJwLibrary)
        if os.path.exists(manifestUser):
            os.remove(manifestUser)

    else:
        dbOriginal = "dbs/userData.db.original"
        dbFromUser = "userBackups/{0}/userData.db".format(telegram_user)
        shutil.copyfile(src=dbOriginal, dst=dbFromUser)

        connection = sqlite3.connect(dbFromUser)
        cursor = connection.cursor()

        cursor.execute("""INSERT INTO Location (LocationId, DocumentId, IssueTagNumber, KeySymbol, Type)
            VALUES (1, ?, ?, "w", 0);""", (documentId, articleId))

        for i in notes:
            cursor.execute("""INSERT INTO InputField ('LocationId', 'TextTag', 'Value') VALUES (1, ?, ?)""",
                           (textareas[i].get("id"), notes[i].replace("'", '"')))

        cursor.execute("UPDATE LastModified SET LastModified = ?", (now_iso,))

        connection.commit()
        connection.close()

        # Calculate hash and last modified date
        hash_digest = calculate_user_data_hash(dbFromUser)
        last_modified_date = get_last_modified_date(dbFromUser)

        # Create manifest data
        manifest_data = {
            "type": 0,
            "name": f"jwlibrary-plus-backup_{now_date}",
            "userDataBackup": {
                "deviceName": "jwlibraryPlus",
                "hash": hash_digest,
                "lastModifiedDate": last_modified_date,
                "databaseName": "userData.db",
                "schemaVersion": schema_version
            },
            "version": 1,
            "creationDate": now.isoformat()
        }

        manifest_file = 'userBackups/{0}/manifest.json'.format(telegram_user)
        with open(manifest_file, 'w', encoding='utf-8') as f:
            json.dump(manifest_data, f, ensure_ascii=False)

        fileName = "userBackups/{0}/jwlibrary-plus-{1}-{2}.jwlibrary".format(telegram_user, documentId, now_date)
        with zipfile.ZipFile(fileName, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.write(dbFromUser, arcname="userData.db")
            zf.write(manifest_file, arcname="manifest.json")
            zf.write(thumbnail_file, arcname="default_thumbnail.png")

        os.remove(dbFromUser)
        os.remove(manifest_file)

    return fileName

########################
### WRITE DOCX AND PDF #
########################

def write_docx_pdf(documentId, title, questions, notes, telegram_user):
    now_date = datetime.now(pytz.timezone('Europe/Madrid')).strftime("%Y-%m-%d")
    document = Document()

    bold_style = document.styles.add_style('Bold List Number', WD_STYLE_TYPE.PARAGRAPH)
    bold_style.font.bold = True

    document.add_heading(title, 0)
    document.add_paragraph('By JW Library Plus - https://github.com/GeiserX/jwlibrary-plus', style="Subtitle")

    for i in range(len(questions)):
        p = document.add_paragraph(style='Bold List Number')
        p.add_run(questions[i].text).font.size = Pt(12)
        document.add_paragraph(notes[i])

    fileNameDoc = "userBackups/{0}/jwlibrary-plus-{1}-{2}.docx".format(telegram_user, documentId, now_date)
    document.save(fileNameDoc)

    fileNamePDF = "userBackups/{0}/jwlibrary-plus-{1}-{2}.pdf".format(telegram_user, documentId, now_date)
    cmd_str = "xvfb-run abiword --to=pdf --to-name='{0}' '{1}'".format(fileNamePDF, fileNameDoc)
    subprocess.run(cmd_str, shell=True)
    return fileNameDoc, fileNamePDF

################
### MAIN CODE ##
################

def main(url, telegram_user, qs_user, language):
    # Set up translation function
    domain = "jwlibraryplus"
    locale_dir = os.path.join(os.path.dirname(__file__), '../locales')
    translation = gettext.translation(domain, localedir=locale_dir, languages=[language], fallback=True)
    trans = translation.gettext

    title, base_text, song, summary, questions, documentId, articleId, q_map, textareas = w_extract_html(url, get_all=True)
    notes = w_query_openai(title, base_text, song, summary, q_map, qs_user, trans)
    filenamejw = write_jwlibrary(documentId, articleId, title, questions, notes, telegram_user, textareas)
    filenamedoc, filenamepdf = write_docx_pdf(documentId, title, questions, notes, telegram_user)
    return filenamejw, filenamedoc, filenamepdf

if __name__ == "__main__":
    # Example usage
    url = "https://www.jw.org/en/library/magazines/watchtower-study-june-2021/do-not-neglect-the-good-you-should-be-doing/"
    telegram_user = "user_id"  # Replace with actual Telegram user ID
    qs_user = [
        "Una ilustración o ejemplo para explicar algún punto principal del párrafo",
        "Una experiencia en concreto, aportando referencias exactas de jw.org, que esté muy relacionada con el párrafo",
        "Una explicación sobre uno de los textos que aparezcan, que aplique al párrafo. Usa la Biblia de Estudio de los Testigos de Jehová"
    ]
    language = "es"  # Replace with the desired language code, e.g., "en" for English
    main(url, telegram_user, qs_user, language)